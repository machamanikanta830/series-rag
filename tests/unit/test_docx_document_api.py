"""Tests for bounded native DOCX document uploads."""

from collections.abc import Iterator
from io import BytesIO

import pytest
from docx import Document as create_docx
from docx.document import Document as DocxDocument
from fastapi.testclient import TestClient

from app.api.dependencies import (
    get_document_catalog,
    get_ingestion_service,
    reset_development_application_state,
)
from app.api.main import MAX_UPLOAD_BYTES, app
from app.models import Document
from app.services.ingestion import IngestionService, IngestionStatistics

client = TestClient(app)


class StubIngestionService(IngestionService):
    """Return fixed statistics while recording parsed DOCX documents."""

    def __init__(self) -> None:
        self.documents: list[Document] = []

    def ingest(self, document: Document) -> IngestionStatistics:
        """Record the parser output and return deterministic statistics."""
        self.documents.append(document)
        return IngestionStatistics(
            document_id=document.document_id,
            chunks_created=len(document.sections),
            embedding_dimension=3,
            vector_store_name="TestVectorStore",
        )


@pytest.fixture(autouse=True)
def isolate_application_state() -> Iterator[None]:
    """Reset shared storage and dependency overrides around every test."""
    reset_development_application_state()
    app.dependency_overrides.clear()
    yield
    app.dependency_overrides.clear()
    reset_development_application_state()


def _docx_bytes(document: DocxDocument) -> bytes:
    """Save a generated DOCX entirely in memory for upload tests."""
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _paragraph_docx(text: str) -> bytes:
    """Create a one-paragraph native DOCX."""
    document = create_docx()
    document.add_paragraph(text)
    return _docx_bytes(document)


def test_native_docx_upload_returns_serialized_ingestion_response() -> None:
    """The API invokes the DOCX parser before the existing ingestion boundary."""
    service = StubIngestionService()
    app.dependency_overrides[get_ingestion_service] = lambda: service

    response = client.post(
        "/documents",
        files={
            "file": (
                "course.docx",
                _paragraph_docx("Native Word lesson."),
                "application/octet-stream",
            )
        },
    )

    assert response.status_code == 201
    [document] = service.documents
    assert response.json() == {
        "document_id": document.document_id,
        "filename": "course.docx",
        "chunks_created": 1,
        "embedding_dimension": 3,
        "vector_store_name": "TestVectorStore",
    }
    assert document.source_name == "course.docx"
    assert document.metadata["filename"] == "course.docx"
    assert document.sections[0].metadata == {"section_type": "paragraph"}


def test_uploaded_docx_chunks_preserve_structural_metadata() -> None:
    """The shared catalog stores heading, table, and paragraph provenance."""
    source = create_docx()
    source.add_heading("Team", level=1)
    table = source.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Alice"
    table.cell(0, 1).text = "Engineer"
    source.add_paragraph("The team supports cloud security.")

    upload_response = client.post(
        "/documents",
        files={
            "file": (
                "team.docx",
                _docx_bytes(source),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    document_id = upload_response.json()["document_id"]
    catalog_document = get_document_catalog().get_document(document_id)
    detail_response = client.get(f"/documents/{document_id}")

    assert upload_response.status_code == 201
    assert catalog_document is not None
    assert [chunk.metadata["section_type"] for chunk in catalog_document.chunks] == [
        "heading",
        "table",
        "paragraph",
    ]
    assert all(
        chunk.metadata["filename"] == "team.docx" for chunk in catalog_document.chunks
    )
    assert [
        chunk["metadata"]["section_type"] for chunk in detail_response.json()["chunks"]
    ] == ["heading", "table", "paragraph"]


def test_empty_docx_returns_stable_unprocessable_response() -> None:
    """A valid package without usable text maps to HTTP 422."""
    response = client.post(
        "/documents",
        files={"file": ("empty.docx", _docx_bytes(create_docx()), "application/docx")},
    )

    assert response.status_code == 422
    assert response.json() == {"detail": "Uploaded DOCX contains no extractable text."}


def test_malformed_docx_returns_stable_unprocessable_response() -> None:
    """Unreadable package bytes are rejected before ingestion."""
    response = client.post(
        "/documents",
        files={"file": ("broken.docx", b"not a DOCX", "application/docx")},
    )

    assert response.status_code == 422
    assert response.json() == {"detail": "Uploaded DOCX is malformed or unreadable."}


def test_oversized_docx_is_rejected_before_parsing() -> None:
    """DOCX uploads retain the existing one-megabyte bounded-read limit."""
    response = client.post(
        "/documents",
        files={
            "file": (
                "large.docx",
                b"x" * (MAX_UPLOAD_BYTES + 1),
                "application/docx",
            )
        },
    )

    assert response.status_code == 413
    assert response.json() == {"detail": "Uploaded file exceeds the 1 MB size limit."}


def test_uploaded_docx_content_is_queryable_through_the_shared_store() -> None:
    """Native DOCX text follows the existing retrieval path unchanged."""
    upload_response = client.post(
        "/documents",
        files={
            "file": (
                "nutrition.docx",
                _paragraph_docx("Oranges contain vitamin C and support nutrition."),
                "application/docx",
            )
        },
    )
    document_id = upload_response.json()["document_id"]

    query_response = client.post(
        "/query",
        json={"question": "What do oranges contain for nutrition?", "top_k": 1},
    )

    assert upload_response.status_code == 201
    assert query_response.status_code == 200
    assert query_response.json()["sources"][0]["document_id"] == document_id
    assert query_response.json()["sources"][0]["source_name"] == "nutrition.docx"


def test_dependency_override_isolation_restores_default_docx_ingestion() -> None:
    """A stub service cannot leak into later DOCX API requests."""
    assert get_ingestion_service not in app.dependency_overrides

    response = client.post(
        "/documents",
        files={
            "file": (
                "default.docx",
                _paragraph_docx("Default DOCX text."),
                "application/docx",
            )
        },
    )

    assert response.status_code == 201
    assert response.json()["vector_store_name"] == "InMemoryVectorStore"
