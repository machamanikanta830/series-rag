"""Tests for ordered native DOCX text and structure extraction."""

from io import BytesIO

import pytest
from docx import Document as create_docx
from docx.document import Document as DocxDocument

from app.chunking import chunk_document
from app.parsers.docx import (
    DOCXNoExtractableTextError,
    DOCXParsingError,
    parse_docx,
)


def _docx_bytes(document: DocxDocument) -> bytes:
    """Save a generated DOCX to in-memory bytes."""
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_extracts_headings_and_multiple_paragraphs_in_order() -> None:
    """Headings and paragraphs become ordered sections with explicit types."""
    source = create_docx()
    source.add_heading("Cloud Security", level=1)
    source.add_paragraph("AWS protects physical infrastructure.")
    source.add_paragraph("Customers protect their data and identities.")

    document = parse_docx(_docx_bytes(source), "course.docx")

    assert [section.text for section in document.sections] == [
        "Cloud Security",
        "AWS protects physical infrastructure.",
        "Customers protect their data and identities.",
    ]
    assert [section.metadata["section_type"] for section in document.sections] == [
        "heading",
        "paragraph",
        "paragraph",
    ]
    assert document.sections[0].metadata["heading_style"] == "Heading 1"
    assert document.metadata == {
        "filename": "course.docx",
        "source_type": "docx",
    }


def test_extracts_table_as_deterministic_pipe_separated_text() -> None:
    """Table cells retain row and column readability without HTML or Markdown."""
    source = create_docx()
    table = source.add_table(rows=3, cols=3)
    values = [
        ["Name", "Role", "Location"],
        ["Alice", "Engineer", "Iowa"],
        ["Bob", "Analyst", "Chicago"],
    ]
    for row, row_values in zip(table.rows, values, strict=True):
        for cell, value in zip(row.cells, row_values, strict=True):
            cell.text = value

    document = parse_docx(_docx_bytes(source), "people.docx")

    [section] = document.sections
    assert section.text == (
        "Name | Role | Location\nAlice | Engineer | Iowa\nBob | Analyst | Chicago"
    )
    assert section.metadata == {"section_type": "table"}


def test_preserves_paragraph_and_table_reading_order() -> None:
    """Body blocks remain in the same order a reader encounters them."""
    source = create_docx()
    source.add_paragraph("Before table.")
    table = source.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Key"
    table.cell(0, 1).text = "Value"
    source.add_paragraph("After table.")

    document = parse_docx(_docx_bytes(source), "ordered.docx")

    assert [section.text for section in document.sections] == [
        "Before table.",
        "Key | Value",
        "After table.",
    ]
    assert [section.metadata["section_type"] for section in document.sections] == [
        "paragraph",
        "table",
        "paragraph",
    ]


def test_chunking_preserves_docx_section_provenance() -> None:
    """Chunk metadata combines filename and structural section type."""
    source = create_docx()
    source.add_heading("Responsibilities", level=2)
    source.add_paragraph("Customers configure access permissions.")
    document = parse_docx(_docx_bytes(source), "security.docx")

    chunks = chunk_document(document, chunk_size=20, chunk_overlap=2)

    assert [chunk.metadata["section_type"] for chunk in chunks] == [
        "heading",
        "paragraph",
    ]
    assert all(chunk.metadata["filename"] == "security.docx" for chunk in chunks)
    assert all("page_number" not in chunk.metadata for chunk in chunks)


def test_rejects_an_empty_docx() -> None:
    """A valid package without supported text does not enter ingestion."""
    with pytest.raises(DOCXNoExtractableTextError, match="no extractable"):
        parse_docx(_docx_bytes(create_docx()), "empty.docx")


def test_rejects_a_malformed_docx() -> None:
    """Invalid ZIP/package bytes become one stable parser-level error."""
    with pytest.raises(DOCXParsingError, match="malformed or unreadable"):
        parse_docx(b"not a DOCX package", "broken.docx")
